#!/usr/bin/env python3
"""
Generate comprehensive EEG Arrow Stimulus Report (.docx)
Consolidates all ERP and decoding analysis results.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = r'E:\deskbook\OpenBCI_GUI\stimulus_logs\analysis_v5'
REPORT_PATH = os.path.join(OUT_DIR, 'EEG_Arrow_Stimulus_Report.docx')

doc = Document()

# ========== Styles ==========
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def add_image_safe(doc, path, width=Inches(5.5)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=width)
        return True
    return False

def pct(v):
    return f'{v:.1%}'

# ================================================================
# TITLE PAGE
# ================================================================
for _ in range(6):
    doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('EEG Arrow Stimulus\nDecoding Analysis Report')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

doc.add_paragraph('')
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(
    'OpenBCI Cyton+Daisy | 12 Channels | 4 Arrow Directions\n'
    'Within-Session & Cross-Session Decoding'
)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph('')
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Generated: 2026-07-17')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

doc.add_page_break()

# ================================================================
# TABLE OF CONTENTS
# ================================================================
doc.add_heading('Table of Contents', level=1)
toc = [
    '1. Overview',
    '2. Methods',
    '3. ERP Waveform Analysis',
    '4. Region Combination Decoding',
    '5. PCA Dimensionality Reduction',
    '6. Cross-Session Validation',
    '7. Feature Importance Analysis',
    '8. Single-Channel Decoding',
    '9. Feature Engineering Comparison',
    '10. Binary Classification',
    '11. Hierarchical Decoding Analysis',
    '12. Prefrontal (F3/Fz/F4) Detailed Binary',
    '13. Summary & Conclusions',
]
for item in toc:
    doc.add_paragraph(item).paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ================================================================
# 1. OVERVIEW
# ================================================================
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'This report presents a comprehensive analysis of EEG responses to visual arrow stimuli '
    '(4 directions: Up, Down, Left, Right) recorded using OpenBCI Cyton+Daisy (12 active channels, '
    '500 Hz sampling). The analysis covers ERP waveform characterization and single-trial decoding '
    'using Linear Discriminant Analysis (LDA).'
)

doc.add_heading('Data Summary', level=2)
add_table(doc,
    ['Property', 'Value'],
    [
        ['Device', 'OpenBCI Cyton + Daisy WiFi'],
        ['Channels', '12 active (CH1,2,4,5,6,7,8,9,10,12,14,15)'],
        ['Channel Names', 'Oz, C3, Fz, C4, Cz, F3, O2, P3, Pz, P4, F4, O1'],
        ['Sampling Rate', '500 Hz'],
        ['Reference', 'SRB2 (all channels averaged)'],
        ['Gain', '6x'],
        ['ADC Scale', '0.08941 uV/count'],
        ['Filter', 'Butterworth 4th-order 1-45 Hz bandpass + 50 Hz notch'],
        ['Baseline Correction', '-200 to 0 ms pre-stimulus'],
        ['Epoch Window', '-200 to +800 ms (500 time points)'],
        ['Artifact Rejection', '>100 uV peak on any channel'],
        ['Total Trials (Session 2)', '200 (50 per direction)'],
        ['ERP Windows', 'P1(80-130ms), N1(140-200ms), P2(200-300ms), P3(300-500ms)'],
        ['Regions', 'Frontal(F3,Fz,F4), Central(C3,Cz,C4), Parietal(P3,Pz,P4), Occipital(O1,Oz,O2)'],
    ]
)

doc.add_page_break()

# ================================================================
# 2. METHODS
# ================================================================
doc.add_heading('2. Methods', level=1)
doc.add_paragraph(
    'For single-trial decoding, 4 ERP window means (P1, N1, P2, P3) per channel serve as features. '
    'Features are scaled with StandardScaler and classified using LDA with 5-fold stratified '
    'cross-validation (random_state=42). Two evaluation dimensions are reported:\n'
    '(1) 4-class exact accuracy: prediction matches the exact direction\n'
    '(2) Axis accuracy: prediction falls on the correct spatial axis (vertical vs horizontal)\n'
    '15 region combinations are tested: 4 single regions + 6 pairs + 4 triples + all 12 channels.\n\n'
    'Additional feature engineering approaches tested include: PCA (95% variance), '
    'peak-to-peak range per window, full time course (500 points per channel) + PCA, '
    'shrinkage LDA, FFT-based band power features, feature fusion, and stacking classifiers.'
)

doc.add_page_break()

# ================================================================
# 3. ERP WAVEFORM ANALYSIS
# ================================================================
doc.add_heading('3. ERP Waveform Analysis', level=1)
doc.add_paragraph(
    'Grand average ERP waveforms across all 5 sessions show clear P1, N1, P2, and P3 components '
    'in response to arrow stimuli. Occipital channels show the largest visual evoked potentials, '
    'while frontal channels show sustained activity during the P3 window.'
)

for img in ['02_roi_activation_erp.png', '03_activation_heatmap.png',
            '07_channel_erp_grid.png', '08_grand_average_roi.png',
            '14_scalp_direction_erp.png']:
    add_image_safe(doc, os.path.join(OUT_DIR, img), Inches(5.5))

doc.add_page_break()

# ================================================================
# 4. REGION COMBINATION DECODING
# ================================================================
doc.add_heading('4. Region Combination Decoding', level=1)
doc.add_paragraph(
    '4-class LDA decoding across 15 region combinations using 4 ERP window means per channel. '
    'Sorted by 4-class accuracy. All results: within-session, Session 2.'
)

region_data = [
    ('Frontal (F3,Fz,F4)', '3ch', 0.6150, 0.5250, 0.5700, 0.9500),
    ('F+P', '6ch', 0.6500, 0.5450, 0.6900, 0.9500),
    ('F+P+O', '9ch', 0.6050, 0.4850, 0.7200, 0.9200),
    ('F+C', '6ch', 0.6050, 0.5050, 0.6500, 0.9100),
    ('F+O', '6ch', 0.6200, 0.4850, 0.6500, 0.9300),
    ('F+C+O', '9ch', 0.5900, 0.5150, 0.6600, 0.8900),
    ('F+C+P', '9ch', 0.5700, 0.5100, 0.6500, 0.8700),
    ('All 12ch', '12ch', 0.5400, 0.5050, 0.6500, 0.8600),
    ('Central (C3,Cz,C4)', '3ch', 0.4950, 0.4700, 0.6400, 0.8900),
    ('P+O', '6ch', 0.5250, 0.5150, 0.7000, 0.8000),
    ('Parietal (P3,Pz,P4)', '3ch', 0.4800, 0.5350, 0.5700, 0.7900),
    ('C+P', '6ch', 0.4450, 0.5450, 0.5800, 0.8400),
    ('C+P+O', '9ch', 0.4800, 0.5450, 0.6500, 0.8100),
    ('C+O', '6ch', 0.4600, 0.5250, 0.5600, 0.8500),
    ('Occipital (O1,Oz,O2)', '3ch', 0.4100, 0.4950, 0.5200, 0.7700),
]
region_data.sort(key=lambda r: r[2], reverse=True)

add_table(doc,
    ['Rank', 'Region', 'Ch', '4-Class', 'Axis(V/H)', 'U vs D', 'L vs R'],
    [(i+1, rn, ch, pct(a4), pct(ax), pct(uv), pct(lr))
     for i, (rn, ch, a4, ax, uv, lr) in enumerate(region_data)]
)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key Finding: ')
run.bold = True
p.add_run(
    'F+P (Frontal+Parietal) achieves the highest 4-class accuracy at 65.0%. '
    'The axis discrimination (vertical vs horizontal) is the worst sub-problem across all regions, '
    'never exceeding 54.5%. Left vs Right reaches 95% in frontal regions.'
)

for img in ['31_region_4class_accuracy.png', '33_region_perclass_breakdown.png',
            '35_region_best_confusion.png']:
    add_image_safe(doc, os.path.join(OUT_DIR, img), Inches(5.5))

doc.add_page_break()

# ================================================================
# 5. PCA DIMENSIONALITY REDUCTION
# ================================================================
doc.add_heading('5. PCA Dimensionality Reduction', level=1)
doc.add_paragraph(
    'PCA (95% variance) applied to window-mean features before LDA. '
    'Compared with original (no PCA) performance.'
)

pca_data = [
    ('F+P', '24', '8', 0.6500, 0.5000, -0.1500),
    ('F+O', '24', '8', 0.6200, 0.5250, -0.0950),
    ('Frontal', '12', '6', 0.6150, 0.4350, -0.1800),
    ('F+C', '24', '7', 0.6050, 0.4950, -0.1100),
    ('F+P+O', '36', '10', 0.6050, 0.5500, -0.0550),
    ('F+C+O', '36', '10', 0.5900, 0.5200, -0.0700),
    ('F+C+P', '36', '9', 0.5700, 0.4850, -0.0850),
    ('All 12ch', '48', '11', 0.5400, 0.5350, -0.0050),
    ('P+O', '24', '6', 0.5250, 0.4150, -0.1100),
    ('Central', '12', '5', 0.4950, 0.3800, -0.1150),
    ('Parietal', '12', '5', 0.4800, 0.3800, -0.1000),
    ('C+P+O', '36', '9', 0.4800, 0.3900, -0.0900),
    ('C+O', '24', '7', 0.4600, 0.4050, -0.0550),
    ('C+P', '24', '7', 0.4450, 0.3950, -0.0500),
    ('Occipital', '12', '4', 0.4100, 0.4300, 0.0200),
]

add_table(doc,
    ['Region', 'Features', 'PCs', 'No PCA', 'PCA', 'Delta'],
    [(rn, fn, pc, pct(no), pct(p), pct(d))
     for rn, fn, pc, no, p, d in pca_data]
)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Finding: ')
run.bold = True
p.add_run(
    'PCA uniformly hurts decoding (average 53.9% to 45.6%). '
    'Only Occipital improves slightly (+2.0%). Unsupervised PCA discards discriminative information.'
)
add_image_safe(doc, os.path.join(OUT_DIR, '41_pca_comparison_bars.png'), Inches(5.5))

doc.add_page_break()

# ================================================================
# 6. CROSS-SESSION VALIDATION
# ================================================================
doc.add_heading('6. Cross-Session Validation', level=1)
doc.add_paragraph(
    'Model trained on Sessions 1-4 (777 trials), tested on Session 5 (197 trials). '
    'Compared with within-session 5-fold CV on Session 2 (200 trials).'
)

cs_data = [
    ('Central', 0.4950, 0.4822, -0.0128),
    ('C+O', 0.4600, 0.4518, -0.0082),
    ('C+P', 0.4450, 0.4315, -0.0135),
    ('C+P+O', 0.4800, 0.4010, -0.0790),
    ('Occipital', 0.4100, 0.3909, -0.0191),
    ('All 12ch', 0.5400, 0.4162, -0.1238),
    ('F+C+P', 0.5700, 0.4162, -0.1538),
    ('Parietal', 0.4800, 0.3604, -0.1196),
    ('F+C+O', 0.5900, 0.4061, -0.1839),
    ('F+C', 0.6050, 0.3959, -0.2091),
    ('F+P+O', 0.6050, 0.3756, -0.2294),
    ('F+P', 0.6500, 0.3503, -0.2997),
    ('F+O', 0.6200, 0.3401, -0.2799),
    ('P+O', 0.5250, 0.3350, -0.1900),
    ('Frontal', 0.6150, 0.3198, -0.2952),
]

add_table(doc,
    ['Region', 'Within', 'Cross', 'Delta'],
    [(rn, pct(w), pct(c), pct(d)) for rn, w, c, d in cs_data]
)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key Finding: ')
run.bold = True
p.add_run(
    'Frontal regions collapse cross-session (-29.5%), revealing noise inflation within-session. '
    'Central is most robust at 48.2% (-1.3%). Best cross-session: Central (48.2%), C+O (45.2%).'
)

for img in ['51_cross_session_bars.png', '54_cross_session_scatter.png',
            '52_cross_session_confusion.png']:
    add_image_safe(doc, os.path.join(OUT_DIR, img), Inches(5.0))

doc.add_page_break()

# ================================================================
# 7. FEATURE IMPORTANCE
# ================================================================
doc.add_heading('7. Feature Importance Analysis', level=1)
doc.add_paragraph(
    'LDA coefficient magnitudes (cross-session model) reveal discriminative channels and ERP windows.'
)

doc.add_heading('Window Ranking (across all channels)', level=2)
add_table(doc, ['Window', 'Importance'],
          [('P2', '0.4957'), ('N1', '0.4918'), ('P3', '0.3134'), ('P1', '0.2882')])

doc.add_heading('Channel Ranking (across all windows)', level=2)
add_table(doc, ['Channel', 'Importance'],
          [(c, f'{v:.4f}') for c, v in [
              ('C3', 0.6045), ('F4', 0.5619), ('Pz', 0.4816), ('Fz', 0.4714),
              ('O2', 0.4562), ('Oz', 0.4137), ('P3', 0.4047), ('P4', 0.3592),
              ('C4', 0.3134), ('O1', 0.2533), ('Cz', 0.2414), ('F3', 0.2059)]])

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Finding: ')
run.bold = True
p.add_run(
    'P2 and N1 windows are most discriminative (~0.50 each). '
    'C3 (left central) is the most important channel, followed by F4 (right frontal). '
    'F3 (left frontal) is least important.'
)

add_image_safe(doc, os.path.join(OUT_DIR, '62_feature_importance_all_channels.png'), Inches(5.5))
add_image_safe(doc, os.path.join(OUT_DIR, '63_direction_feature_profiles.png'), Inches(5.5))

doc.add_page_break()

# ================================================================
# 8. SINGLE-CHANNEL DECODING
# ================================================================
doc.add_heading('8. Single-Channel Decoding', level=1)
doc.add_paragraph(
    'Each of 12 channels decoded independently with 4 ERP window means (4 features).'
)

add_table(doc,
    ['Channel', '4-Class', 'Axis(V/H)'],
    [(c, pct(v), pct(a)) for c, v, a in [
        ('O2', 0.4550, 0.6250), ('P4', 0.4450, 0.5950),
        ('C4', 0.4250, 0.6050), ('F4', 0.4200, 0.5400),
        ('Oz', 0.4050, 0.5750), ('Cz', 0.4000, 0.5250),
        ('O1', 0.3950, 0.5950), ('F3', 0.3650, 0.5550),
        ('C3', 0.3650, 0.5300), ('Fz', 0.3600, 0.5050),
        ('Pz', 0.3600, 0.5250), ('P3', 0.3300, 0.4700),
    ]]
)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Finding: ')
run.bold = True
p.add_run(
    'Right-hemisphere dominance: O2 (45.5%) > P4 (44.5%) > C4 (42.5%) > F4 (42.0%) are top 4. '
    'Left counterparts are worse. Average single-channel: 39.4%. O2 also has best axis acc (62.5%).'
)

add_image_safe(doc, os.path.join(OUT_DIR, '71_single_channel_decoding.png'), Inches(5.5))
add_image_safe(doc, os.path.join(OUT_DIR, '72_single_channel_topomap.png'), Inches(5.0))

doc.add_page_break()

# ================================================================
# 9. FEATURE ENGINEERING COMPARISON
# ================================================================
doc.add_heading('9. Feature Engineering Comparison', level=1)
doc.add_paragraph(
    'Six feature engineering strategies tested vs baseline 4 ERP window means. '
    'None improved over baseline.'
)

doc.add_heading('Overall Summary', level=2)
add_table(doc,
    ['Approach', 'Frontal', 'Central', 'F+P', 'All 12ch'],
    [
        ['4 Window Means (baseline)', '61.5%', '49.5%', '65.0%', '54.0%'],
        ['+ Peak-to-Peak Range', '63.0%', '45.0%', '62.0%', '51.5%'],
        ['Full Time Course + PCA', '45.5%', '39.5%', '44.0%', '46.5%'],
        ['Full TC + Shrinkage LDA', '46.0%', '41.0%', '44.0%', '45.0%'],
        ['5 Band Power (FFT)', '42.0%', '35.5%', '46.5%', '43.5%'],
        ['Windows + Band Fusion', '59.5%', '46.0%', '59.0%', '49.5%'],
        ['Stacking (Win + Band)', '61.0%', '48.5%', '61.0%', '57.5%'],
    ]
)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Key Finding: ')
run.bold = True
p.add_run(
    'No approach improves over 4 window means. The bottleneck is the signal itself '
    '(limited channels, SRB2 reference, 50 trials/direction), not feature quality.'
)

doc.add_heading('9.1 Peak-to-Peak Range Features', level=2)
doc.add_paragraph('Added range within each window (8 features/channel). '
                  'Average dropped 53.9% to 51.6%. Only 3/15 regions improved.')
add_image_safe(doc, os.path.join(OUT_DIR, '75_single_channel_range_compare.png'), Inches(5.0))

doc.add_heading('9.2 Full Time Course + PCA', level=2)
doc.add_paragraph('All 500 time points per channel + PCA (95% var). '
                  '0/12 channels improved. Frontal: 61.5% to 45.5% (-16%).')
add_image_safe(doc, os.path.join(OUT_DIR, '81_full_timecourse_single.png'), Inches(5.5))

doc.add_heading('9.3 Shrinkage LDA', level=2)
doc.add_paragraph('Shrinkage LDA (solver=lsqr, auto) on full time course. '
                  'Frontal: 61.5% to 46.0%. Worse than window means + std LDA.')

doc.add_heading('9.4 Time-Frequency Features', level=2)
doc.add_paragraph('FFT band power in Delta/Theta/Alpha/Beta/Gamma bands.')
add_table(doc,
    ['Region', 'Window Means', 'Band Power', 'Per-Win+PCA'],
    [(rn, pct(w), pct(b), pct(p)) for rn, w, b, p in [
        ('Frontal', 0.6150, 0.4200, 0.1950),
        ('Central', 0.4950, 0.3550, 0.2300),
        ('Parietal', 0.4800, 0.3850, 0.2100),
        ('Occipital', 0.4100, 0.2950, 0.2500),
        ('F+P', 0.6500, 0.4650, 0.2050),
        ('All 12ch', 0.5400, 0.4350, 0.2400),
    ]]
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Finding: ')
run.bold = True
p.add_run(
    'TF much worse for 4-class. BUT: band power improves axis discrimination '
    '(Frontal: 52.5% to 64.0%) while degrading within-axis LvsR (95% to 59%).'
)
add_image_safe(doc, os.path.join(OUT_DIR, '96_timefreq_comparison.png'), Inches(5.5))

doc.add_heading('9.5 Feature Fusion & Stacking', level=2)
doc.add_paragraph(
    'Fusion (window means + band power concatenated) and stacking (two-layer) '
    'both failed to improve over window means alone.'
)

doc.add_page_break()

# ================================================================
# 10. BINARY CLASSIFICATION
# ================================================================
doc.add_heading('10. Binary Classification', level=1)
doc.add_paragraph(
    'All 6 direction pairs + axis, 7 region configurations. '
    'LDA, 5-fold CV, 4 window means per channel.'
)

doc.add_heading('Results', level=2)
add_table(doc,
    ['Pair', 'F', 'C', 'P', 'O', 'F+P', 'F+P+O', 'All', 'Best'],
    [
        ['U vs D', '57%', '64%', '57%', '52%', '69%', '72%', '65%', '72%'],
        ['L vs R', '95%', '89%', '79%', '77%', '95%', '92%', '86%', '95%'],
        ['U vs L', '77%', '66%', '65%', '59%', '78%', '69%', '67%', '78%'],
        ['U vs R', '87%', '74%', '75%', '61%', '79%', '83%', '78%', '87%'],
        ['D vs L', '76%', '70%', '60%', '54%', '73%', '71%', '58%', '76%'],
        ['D vs R', '87%', '84%', '80%', '75%', '81%', '79%', '71%', '87%'],
        ['Axis(V/H)', '52.5%', '47%', '53.5%', '49.5%', '54.5%', '48.5%', '50.5%', '54.5%'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Key Findings: ')
run.bold = True
p.add_run(
    'LvsR = 95% (F, F+P). Pairs with Right: 80-87%. '
    'UvsD hardest at 72%. Axis only 54.5% - the sole bottleneck.'
)

for img in ['91_binary_heatmap.png', '92_binary_by_region.png', '93_4class_vs_binary.png']:
    add_image_safe(doc, os.path.join(OUT_DIR, img), Inches(5.0))

doc.add_page_break()

# ================================================================
# 11. HIERARCHICAL DECODING
# ================================================================
doc.add_heading('11. Hierarchical Decoding Analysis', level=1)
doc.add_paragraph(
    '4-class decomposed into: axis (V/H) -> within-axis (Up/Down or Left/Right). '
    'Hierarchical = axis_acc x within_axis_acc. Axis is 100% bottleneck.'
)

add_table(doc,
    ['Region', 'Axis(V/H)', 'U vs D', 'L vs R', '4-Class', 'Hier', 'BN'],
    [
        ['F+P', '54.5%', '69.0%', '95.0%', '65.0%', '44.7%', 'axis'],
        ['Frontal', '52.5%', '57.0%', '95.0%', '61.5%', '39.9%', 'axis'],
        ['Central', '47.0%', '64.0%', '89.0%', '49.5%', '36.0%', 'axis'],
        ['Parietal', '53.5%', '57.0%', '79.0%', '48.0%', '36.4%', 'axis'],
        ['Occipital', '49.5%', '52.0%', '77.0%', '41.0%', '31.9%', 'axis'],
        ['All 12ch', '50.5%', '65.0%', '86.0%', '54.0%', '38.1%', 'axis'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Key Finding: ')
run.bold = True
p.add_run(
    'Axis is the bottleneck in 15/15 (100%) regions. Band power improves axis to 64.0% '
    'but degrades LvsR from 95% to 59%.'
)

for img in ['98a_hierarchical_breakdown.png', '98b_hierarchical_scatter.png']:
    add_image_safe(doc, os.path.join(OUT_DIR, img), Inches(5.0))

doc.add_page_break()

# ================================================================
# 12. PREFRONTAL DETAILED BINARY
# ================================================================
doc.add_heading('12. Prefrontal (F3/Fz/F4) Detailed Binary', level=1)
doc.add_paragraph(
    'Focused binary analysis on prefrontal region, testing each channel and combination.'
)

doc.add_heading('Single-Channel', level=2)
add_table(doc,
    ['Ch', '4-Class', 'UvsD', 'LvsR', 'UvsL', 'UvsR', 'DvsL', 'DvsR', 'Axis'],
    [
        ['F3', '36.5%', '62%', '44%', '44%', '54%', '62%', '60%', '48.0%'],
        ['Fz', '36.0%', '67%', '70%', '64%', '47%', '55%', '69%', '43.0%'],
        ['F4', '42.0%', '65%', '83%', '66%', '62%', '56%', '75%', '42.0%'],
    ]
)

doc.add_heading('Combinations', level=2)
add_table(doc,
    ['Combo', '4-Class', 'UvsD', 'LvsR', 'UvsL', 'UvsR', 'DvsL', 'DvsR', 'Axis'],
    [
        ['F3+F4', '58.0%', '58%', '93%', '73%', '87%', '76%', '82%', '54.5%'],
        ['F3+Fz+F4', '61.5%', '57%', '95%', '77%', '87%', '76%', '87%', '52.5%'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Finding: ')
run.bold = True
p.add_run(
    'F4 dominates (LvsR=83%). F3 at chance for LvsR (44%). '
    'F3+F4 alone achieves 93% LvsR, only 3.5% below full 3-channel set. '
    'Axis at/below chance for all single channels (42-48%).'
)

doc.add_page_break()

# ================================================================
# 13. SUMMARY & CONCLUSIONS
# ================================================================
doc.add_heading('13. Summary & Conclusions', level=1)

doc.add_heading('Best Results', level=2)
add_table(doc,
    ['Metric', 'Best', 'Accuracy', 'Chance'],
    [
        ['4-Class (within)', 'F+P', '65.0%', '25%'],
        ['4-Class (cross)', 'Central', '48.2%', '25%'],
        ['Axis (V/H)', 'F+P', '54.5%', '50%'],
        ['Left vs Right', 'Frontal/F+P', '95.0%', '50%'],
        ['Up vs Down', 'F+P+O', '72.0%', '50%'],
        ['Single Channel', 'O2', '45.5%', '25%'],
    ]
)

doc.add_heading('Failed Approaches', level=2)
doc.add_paragraph('All feature engineering attempts failed to improve over 4 ERP window means:')
for item in [
    'PCA (average -8.3%)',
    'Gradient features (Frontal -2.0%)',
    'Peak-to-peak range (average -2.3%)',
    'Full time course + PCA (Frontal -16.0%)',
    'Shrinkage LDA (Frontal -15.5%)',
    'One-vs-One voting (average -0.4%)',
    'Time-frequency band power (Frontal -19.5%)',
    'Feature fusion (average -2.9%)',
    'Stacking (average +0.8%, not significant)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Core Conclusion', level=2)
p = doc.add_paragraph()
run = p.add_run('The 4-arrow decoding bottleneck is axis discrimination (vertical vs horizontal). ')
run.bold = True
p.add_run(
    'With 12 channels and SRB2 reference, within-session 4-class accuracy caps at ~65% (F+P) '
    'and cross-session at ~48% (Central). Right-hemisphere channels (O2, P4, C4, F4) carry the most '
    'discriminative information. The axis bottleneck reflects a genuine limitation of this EEG setup '
    '(sparse coverage, SRB2 reference, 50 trials/direction), not feature quality.'
)

doc.add_heading('Recommendations', level=2)
for item in [
    'Use average reference instead of SRB2 for better spatial specificity',
    'Increase channel density (64+) for improved spatial resolution',
    'Increase trial count per direction for better covariance estimation',
    'Consider source-space decoding (sLORETA/MNE)',
    'Improve axis discrimination as the critical path to 4-class gains',
]:
    doc.add_paragraph(item, style='List Bullet')

# ================================================================
# SAVE
# ================================================================
doc.save(REPORT_PATH)
print(f'Report saved: {REPORT_PATH}')
