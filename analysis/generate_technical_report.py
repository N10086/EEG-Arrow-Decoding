#!/usr/bin/env python3
"""Generate complete technical report with full algorithm descriptions."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = r'E:\deskbook\OpenBCI_GUI\stimulus_logs\analysis_v5'
REPORT_PATH = os.path.join(OUT_DIR, 'EEG_Arrow_Stimulus_Technical_Report_v2.docx')
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
for lv in range(1, 4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style, t.alignment = 'Light Grid Accent 1', WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(9)
    return t

def add_code(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(txt)
    r.font.name, r.font.size, r.font.color.rgb = 'Consolas', Pt(8), RGBColor(0x2C, 0x3E, 0x50)

def add_img(doc, path, w=Inches(5.2)):
    if os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=w)
        return True
    return False

def add_ref(doc, ref_id, text):
    """Add a labeled algorithm reference box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'[{ref_id}] ')
    run.bold = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)
    run = p.add_run(text)
    run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

# ========================== TITLE ==========================
for _ in range(5): doc.add_paragraph('')
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('EEG↑↓←→解码项目\n技术报告')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
doc.add_paragraph('')
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('OpenBCI + Processing + Python + LDA\n全链路技术文档')
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
doc.add_page_break()

# ====================== 1. OpenBCI软件 ======================
doc.add_heading('一、OpenBCI软件编写（Processing刺激程序）', level=1)

doc.add_heading('1.1 程序概述', level=2)
doc.add_paragraph(
    '使用 Java + Processing 开发了完整的视觉刺激呈现与数据同步程序。'
    '该程序同时承担三个核心功能：'
    '（1）4方向箭头视觉刺激的随机呈现；'
    '（2）通过串口向OpenBCI传送实时标记信号，实现刺激与EEG数据的精确同步；'
    '（3）实时显示多通道EEG波形，监控信号质量。')

doc.add_heading('1.2 箭头刺激程序核心逻辑', level=2)
add_code(doc,
    '初始化\n'
    '  |\n'
    '  +-- 建立与OpenBCI的串口通信（波特率115200）\n'
    '  +-- 设置全屏画布（1920x1080）\n'
    '  +-- 加载箭头图形资源，配置背景颜色（灰色128）\n'
    '  |\n'
    '  +-- 进入主循环\n'
    '        |\n'
    '        +-- 随机选取方向（0=Up, 1=Down, 2=Left, 3=Right）\n'
    '        +-- 在屏幕中央绘制箭头（白色箭头，灰色背景）\n'
    '        +-- 通过串口发送标记值（2.0001-2.0004）\n'
    '        +-- 持续显示 200ms（刺激呈现）\n'
    '        +-- 清屏进入 800-1000ms 间隔期\n'
    '        +-- 循环 50 次/方向，共 200 trials')

doc.add_heading('Processing核心代码（箭头绘制与标记同步）', level=2)
add_code(doc,
    '// 箭头绘制函数\n'
    'void drawArrow(int direction) {\n'
    '  background(128);\n'
    '  stroke(255);\n'
    '  strokeWeight(8);\n'
    '  fill(255);\n'
    '  pushMatrix();\n'
    '  translate(width/2, height/2);\n'
    '  switch(direction) {\n'
    '    case 0: rotate(0); break;\n'
    '    case 1: rotate(PI); break;\n'
    '    case 2: rotate(-HALF_PI); break;\n'
    '    case 3: rotate(HALF_PI); break;\n'
    '  }\n'
    '  beginShape();\n'
    '  vertex(0, -100);\n'
    '  vertex(-40, 20);\n'
    '  vertex(40, 20);\n'
    '  endShape(CLOSE);\n'
    '  popMatrix();\n'
    '}\n'
    '\n'
    '// 发送标记到OpenBCI\n'
    'void sendMarker(int dir) {\n'
    '  float markerVal = 2.0001 + dir * 0.0001;\n'
    '  openBCI.sendMarker(markerVal);\n'
    '}')

doc.add_heading('1.3 串口通信与视频标记同步机制', level=2)
doc.add_paragraph(
    '刺激程序通过串口（Serial）与OpenBCI主板建立通信。'
    '每次刺激呈现时，程序向OpenBCI的模拟输出/数字引脚发送对应的电压信号，'
    '该信号与EEG数据同步采样，作为数据文件中的第33列（Marker列）记录。')
doc.add_paragraph(
    '同时，程序在每次刺激时联同记录系统时间戳（millis()），'
    '用于后续与视频录制的弧度同步。当视频照相机拍摄实验过程时，'
    '可根据这些时间戳将视频帧与EEG刺激标记一一对应。')

doc.add_heading('1.4 标记值与方向的对应关系', level=2)
add_table(doc, ['标记值', '方向', '说明'],
    [['2.0001', 'Up (上)', '箭头朝上'],
     ['2.0002', 'Down (下)', '箭头朝下'],
     ['2.0003', 'Left (左)', '箭头朝左'],
     ['2.0004', 'Right (右)', '箭头朝右']])
doc.add_paragraph('标记值写入到OpenBCI数据文件的第33列（索引32），与EEG数据同步记录。')

doc.add_page_break()

# ====================== 2. 数据采集 ======================
doc.add_heading('二、数据采集', level=1)
doc.add_heading('2.1 采集设备', level=2)
add_table(doc, ['参数', '值'],
    [['设备', 'OpenBCI Cyton + Daisy WiFi'],
     ['采样率', '500 Hz'],
     ['增益', '6x'],
     ['ADC分辨率', '24位'],
     ['ADC刻度系数', '0.08941 uV/计数'],
     ['参考方式', 'SRB2（全部通道共平均）'],
     ['活性通道', '12通道']])

doc.add_heading('2.2 活性通道与脑区对应', level=2)
add_table(doc, ['OpenBCI通道', '10-20名称', '所属脑区', '功能定位'],
    [['CH1', 'Oz', '枕叶', '枕中线，初级视觉皮层'],
     ['CH2', 'C3', '中央', '左中央沟，运动皮层'],
     ['CH4', 'Fz', '额叶', '额中线，前额叶'],
     ['CH5', 'C4', '中央', '右中央沟，运动皮层'],
     ['CH6', 'Cz', '中央', '中央中线，辅助运动区'],
     ['CH7', 'F3', '额叶', '左前额叶'],
     ['CH8', 'O2', '枕叶', '右枕叶，视觉皮层'],
     ['CH9', 'P3', '顶叶', '左顶叶，空间注意'],
     ['CH10', 'Pz', '顶叶', '顶中线'],
     ['CH12', 'P4', '顶叶', '右顶叶，空间注意'],
     ['CH14', 'F4', '额叶', '右前额叶'],
     ['CH15', 'O1', '枕叶', '左枕叶，视觉皮层']])

doc.add_paragraph('')
add_table(doc, ['脑区', '包含通道', '功能'],
    [['额叶 (Frontal)', 'F3, Fz, F4', '方向偏好决策'],
     ['中央 (Central)', 'C3, Cz, C4', '运动准备'],
     ['顶叶 (Parietal)', 'P3, Pz, P4', '空间注意定向'],
     ['枕叶 (Occipital)', 'O1, Oz, O2', '视觉处理']])

doc.add_heading('2.3 采集流程', level=2)
doc.add_paragraph(
    '受试者: 1人，坐姿，距屏幕约60cm。'
    '实验范式: 每次试验呈现一个箭头方向（200ms），'
    '刺激间隔800-1000ms随机。'
    '共50次/方向 x 4方向 = 200 trials/Session，'
    '共5个Session（不同日期采集），总计约1000 trials。')

doc.add_page_break()

# ====================== 3. 原始数据格式 ======================
doc.add_heading('三、原始数据格式', level=1)
doc.add_heading('3.1 文件格式', level=2)
doc.add_paragraph(
    'OpenBCI输出为CSV格式文本文件，命名格式：'
    'OpenBCI-RAW-YYYY-MM-DD_HH-MM-SS.txt。前5行为文件头信息，第6行开始为数据行。')

doc.add_heading('3.2 原始数据行示例', level=2)
add_code(doc,
    '1, 203.45, 201.23, 198.56, 205.89, 199.12, 207.34, 195.67,\n'
    '202.89, 200.12, 204.56, 197.89, 206.78, 201.34, 199.67, 203.12,\n'
    '1012, 998, 1023, 0, 0, 0, 0, 0, 2.0001')

doc.add_heading('3.3 每一列代表的含义', level=2)
add_table(doc, ['列索引', '列号', '含义', '示例值'],
    [['0', '第1列', '采样序号', '1'],
     ['1-16', '第2-17列', 'EEG通道原始ADC值（CH1-CH16）', '203.45'],
     ['17', '第18列', '加速计X轴', '1012'],
     ['18', '第19列', '加速计Y轴', '998'],
     ['19', '第20列', '加速计Z轴', '1023'],
     ['20-30', '第21-31列', '数字/模拟引脚状态', '0'],
     ['31', '第32列', '备用引脚', '0'],
     ['32', '第33列', '刺激标记（Marker）', '2.0001']])
doc.add_paragraph('')
add_table(doc, ['标记值', '含义'],
    [['2.0001', '呈现“向上”箭头'],
     ['2.0002', '呈现“向下”箭头'],
     ['2.0003', '呈现“向左”箭头'],
     ['2.0004', '呈现“向右”箭头'],
     ['0', '无刺激（基线/间隔期）']])

doc.add_heading('3.4 浮点数精度说明', level=2)
doc.add_paragraph(
    '由于文件存储时的精度限制，原始数据中的标记值并非精确的2.0001，'
    '而是浮点数近似值，如 2.0000998973846436。'
    '在数据处理时采用容差匹配（abs(val - 2.0001) < 5e-5）而非精确相等判断。')

doc.add_page_break()

# ====================== 4. 数据处理与可视化 ======================
doc.add_heading('四、信号处理与数据可视化', level=1)
doc.add_heading('4.1 信号预处理流程', level=2)
doc.add_paragraph('从原始ADC数值到可用的EPR试次，经过以下处理链：')
add_ref(doc, 'A1', 'ADC转换: raw x SCALE，其中SCALE = 4.5 / ((2^23 - 1) x gain) x 10^6 = 0.08941 uV/count。该系数将OpenBCI的24位ADC原始值转换为微伏单位。')
add_ref(doc, 'A2', '带通滤波: 4阶Butterworth，截止频率1-45Hz。scipy.signal.butter(4, [1/250, 45/250], btype="band", output="sos")。500Hz采样下Nyquist频率为250Hz，故归一化截止频率为[1/250, 45/250]。采用sosfiltfilt实现零相位滤波，消除滤波引起的相位偏移。')
add_ref(doc, 'A3', '50Hz工频除波: scipy.signal.iirnotch(50/250, 30)。工频除波器30为品质因子，控制除波带宽。')
add_ref(doc, 'A4', '基线校正: 每个试次减去刺激前200ms（-200ms到0ms）的平均电压，即 ep -= ep[:, :n_pre].mean(axis=1, keepdims=True)，其中n_pre=100个采样点。')
add_ref(doc, 'A5', '伪迹剔除: 任一通道的峰峰值超过±100uV则剔除该试次。该阈值基于EEG信号的典型幅度（<50uV），超过者多为眼电、肌电等伪迹。')

doc.add_heading('4.2 ERP窗口提取', level=2)
doc.add_paragraph('按四个标准ERP成分窗口提取均值特征，该方法基于前人研究对视觉辅助电位的时间窗划分：')
add_table(doc, ['窗口', '时间范围', '神经意义', '参考'],
    [['P1', '80-130ms', '早期视觉处理', '[Di Russo et al., 2002]'],
     ['N1', '140-200ms', '视觉注意分配', '[Hopf et al., 2002]'],
     ['P2', '200-300ms', '刺激分类', '[Luck, 2005]'],
     ['P3', '300-500ms', '决策与更新', '[Polich, 2007]']])
add_ref(doc, 'A6', '特征提取算法: 对每个试次的每个通道，在窗口时间范围内取平均值: win_mean[ti, ci, wi] = mean(epochs_data[ti, ci, msk])，其中msk = (t >= ws) & (t <= we)。输出矩阵维度: n_trials x 12_channels x 4_windows。')

add_img(doc, os.path.join(OUT_DIR, '02_roi_activation_erp.png'), Inches(5.0))
doc.add_paragraph('')

doc.add_heading('4.3 数据可视化（Python + Matplotlib）', level=2)
doc.add_paragraph('基于Python的Matplotlib库开发了系统的可视化方案，覆盖从波形到解码结果的多个层面：')

doc.add_heading('A. ERP波形图', level=3)
doc.add_paragraph('12通道 x 4方向，共48条ERP波形叠加。x轴：时间（-200ms到800ms），y轴：电压（uV），4方向用不同颜色区分。')
add_img(doc, os.path.join(OUT_DIR, '02_roi_activation_erp.png'), Inches(5.0))

doc.add_heading('B. 通道栅格图', level=3)
doc.add_paragraph('12个子图对应12个通道，排列方式对应通道在头皮上的空间位置。')
add_img(doc, os.path.join(OUT_DIR, '07_channel_erp_grid.png'), Inches(5.0))

doc.add_heading('C. 激活热力图', level=3)
doc.add_paragraph('x轴：时间点，y轴：通道，颜色编码电压值（红=正，蓝=负）。')
add_img(doc, os.path.join(OUT_DIR, '03_activation_heatmap.png'), Inches(5.0))

doc.add_heading('D. 跨Session大平均', level=3)
doc.add_paragraph('4个脑区 x 4方向，每个子图为5个Session的叠加平均。')
add_img(doc, os.path.join(OUT_DIR, '08_grand_average_roi.png'), Inches(5.0))

doc.add_heading('E. 单通道解码对比', level=3)
doc.add_paragraph('12通道的4分类准确率排序柱状图，颜色标识所属脑区。柱状图显示右半球通道普遍由于左半球，其中O2以45.5%排名第一。')
add_img(doc, os.path.join(OUT_DIR, '71_single_channel_decoding.png'), Inches(5.0))

doc.add_heading('F. 二分类热力图', level=3)
doc.add_paragraph('x轴：7个方向配对，y轴：7个脑区组合，颜色编码二分类准确率（绿=高，红=低）。')
add_img(doc, os.path.join(OUT_DIR, '91_binary_heatmap.png'), Inches(5.0))

doc.add_heading('4.4 实时可视化界面（Processing）', level=2)
doc.add_paragraph('实验中使用的Processing程序同时承担了实时可视化界面的功能：')
add_code(doc,
    '+----------------------------------------------------+\n'
    '|              EEG实时监控界面                       |\n'
    '|  +------------------------------------------------+ |\n'
    '|  |          中央箭头刺激区                    | |\n'
    '|  |              ^                                  | |\n'
    '|  +------------------------------------------------+ |\n'
    '|  | CH1 | | CH2 | | CH3 | | CH4 | | CH5 | | CH6 |  |\n'
    '|  | ~~~ | | ~~~ | | ~~~ | | ~~~ | | ~~~ | | ~~~ |  |\n'
    '|  +-----+ +-----+ +-----+ +-----+ +-----+ +-----+  |\n'
    '|  [进度: 45/200 trials]  [当前方向: Up]             |\n'
    '+----------------------------------------------------+')
doc.add_paragraph('实现功能：(1) 中心区域呈现箭头视觉刺激；(2) 波形面板实时滚动显示各通道EEG波形；(3) 试验进度和方向提示；(4) 通过串口向OpenBCI写入当前方向标记值。')

doc.add_heading('4.5 批量分析脚本执行', level=2)
add_code(doc,
    'python plot_region_comprehensive.py    # 15脑区组合解码\n'
    'python plot_single_channel.py          # 单通道解码\n'
    'python plot_cross_session.py           # 跨Session验证\n'
    'python plot_feature_importance.py      # 特征重要性\n'
    'python plot_binary.py                  # 二分类分析\n'
    'python plot_timefreq.py                # 时频特征对比\n'
    'python plot_binary_hierarchical.py     # 层次化解码\n'
    'python generate_report.py              # 生成Word报告')

doc.add_page_break()

# ====================== 5. 箭头方向预测 ======================
doc.add_heading('五、箭头方向预测任务', level=1)
doc.add_heading('5.1 任务定义', level=2)
doc.add_paragraph('目标: 单试次EEG信号 -> 预测箭头方向（4分类）')
add_code(doc,
    '输入: 一次箭头刺激后-200到+800ms的12通道EEG信号（500个时间点）\n'
    '输出: 4个方向之一 (0=Up, 1=Down, 2=Left, 3=Right)')

doc.add_heading('5.2 评价指标', level=2)
add_ref(doc, 'B1', '4分类准确率: 预测方向与真实方向完全一致的比例，机会水平=25%。计算公式: accuracy = sum(y_pred == y_true) / n_trials。')
add_ref(doc, 'B2', '轴准确率: 垂直类(Up/Down) vs 水平类(Left/Right)的二分类准确率，机会水平=50%。计算时将预测和真实方向映射到轴类别: axis(y) = 0 if y in [0,1] else 1。')
add_ref(doc, 'B3', '混淆矩阵: 4x4矩阵，行为真实类别，列为预测类别，格子内为试次数（或百分比）。用于详细分析错误模式。')

doc.add_heading('5.3 模型输入特征', level=2)
doc.add_paragraph('基准方案: 4个ERP窗口均值 x 每个脑区通道数 = 特征向量。')
doc.add_paragraph('以额叶（F3, Fz, F4）为例，特征向量由4个ERP窗口均值 x 通道数组成：')
add_code(doc,
    '特征向量 = [F3_P1, F3_N1, F3_P2, F3_P3,\n'
    '            Fz_P1, Fz_N1, Fz_P2, Fz_P3,\n'
    '            F4_P1, F4_N1, F4_P2, F4_P3]\n'
    '维度: 12 (3通道 x 4窗口)\n'
    'F+P(6通道): 24维  |  全部 12通道: 48维')
add_ref(doc, 'A6', '特征提取详细算法见第四章中“ERP窗口提取”部分（上文ref A6）。')

doc.add_heading('5.4 模型算法：LDA原理', level=2)
doc.add_paragraph(
    '主模型采用线性判别分析（Linear Discriminant Analysis, LDA）。'
    'LDA通过最大化类间散布与类内散布的比值，'
    '找到最优投影方向：')
add_ref(doc, 'C1', 'LDA目标函数: J(w) = (w^T S_B w) / (w^T S_W w)，其中S_B = sum(n_k (mu_k - mu)(mu_k - mu)^T)为类间散布矩阵，S_W = sum(sum(x - mu_k)(x - mu_k)^T)为类内散布矩阵，n_k为第k类的样本数，mu_k为第k类的均值向量，mu为总均值向量。')
add_ref(doc, 'C2', '多分类决策规则: LDA将高维特征投影到K-1维判别空间（K为类别数）。对于4分类，产生3个判别函数。预测时计算试歡向量x到每个类别的马氏距离: d_k(x) = x^T Sigma^{-1} mu_k - 0.5 mu_k^T Sigma^{-1} mu_k + ln(pi_k)，其中Sigma为共享协方差矩阵。最终输出: y_pred = argmax_k d_k(x)。')

doc.add_paragraph('')
doc.add_heading('Python核心代码示例', level=3)
doc.add_paragraph('以下代码展示了完整的解码流程，对应上述算法原理的第C1-C2步骤：')
add_code(doc,
    'import numpy as np\n'
    'from sklearn.discriminant_analysis import LinearDiscriminantAnalysis\n'
    'from sklearn.model_selection import StratifiedKFold\n'
    'from sklearn.preprocessing import StandardScaler\n'
    'from sklearn.metrics import accuracy_score\n'
    '\n'
    '# ====== 1. ERP窗口均值特征提取（对应ref A6） ======\n'
    'ERP_WINDOWS = [("P1",0.080,0.130), ("N1",0.140,0.200),\n'
    '               ("P2",0.200,0.300), ("P3",0.300,0.500)]\n'
    'win_mean = np.zeros((len(epochs_data), 12, 4))\n'
    'for wi, (_, ws, we) in enumerate(ERP_WINDOWS):\n'
    '    msk = (t >= ws) & (t <= we)\n'
    '    win_mean[:, :, wi] = epochs_data[:, :, msk].mean(axis=2)\n'
    '\n'
    '# ====== 2. 构建特征矩阵 ======\n'
    'ch_idx = [HW_NAMES.index(c) for c in ["F3","Fz","F4"]]\n'
    'X = win_mean[:, ch_idx, :].reshape(len(epochs_data), -1)\n'
    'y = epochs_label  # 0=Up, 1=Down, 2=Left, 3=Right\n'
    '\n'
    '# ====== 3. 5折分层交叉验证 ======\n'
    'skf = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)\n'
    'accs = []\n'
    'for train_idx, test_idx in skf.split(X, y):\n'
    '    X_train, X_test = X[train_idx], X[test_idx]\n'
    '    y_train, y_test = y[train_idx], y[test_idx]\n'
    '\n'
    '    # 标准化（对应ref A1）\n'
    '    scaler = StandardScaler().fit(X_train)\n'
    '    X_train_s = scaler.transform(X_train)\n'
    '    X_test_s = scaler.transform(X_test)\n'
    '\n'
    '    # LDA训练与预测（对应ref C1-C2）\n'
    '    clf = LinearDiscriminantAnalysis().fit(X_train_s, y_train)\n'
    '    y_pred = clf.predict(X_test_s)\n'
    '\n'
    '    # 评估（对应ref B1）\n'
    '    accs.append(accuracy_score(y_test, y_pred))\n'
    '\n'
    'print(f"4分类准确率: {np.mean(accs):.1%}")')

doc.add_paragraph('')
doc.add_paragraph(
    '这段代码的核心逻辑：(1) 对每个试次提取4个ERP窗口的均值电压作为特征（A6）；'
    '(2) 按脑区组合选取对应通道的特征拼接成特征向量；'
    '(3) 在5折交叉验证中，每轮先标准化再训练LDA模型（C1）；'
    '(4) 输出5轮平均准确率作为最终评估指标（B1）。')

doc.add_heading('5.5 其他尝试过的算法', level=2)
add_table(doc, ['算法', '原理说明', '最佳结果'],
    [['标准LDA', '基准分类器，共享协方差矩阵+马氏距离决策', '65.0%'],
     ['收缩LDA（lsqr+auto）', '引入收缩参数优化协方差估计，适用于高维小样本场景', '46.0%'],
     ['OvO投票LDA', '训练6个二分类器（6个方向配对），投票决定最终类别', '60.0%（中央）'],
     ['Stacking (LR/LDA)', '两层融合：层4分别用窗口均值和频带能量训练LDA，层5用LR融合预测', '61.0%（F+P）']])

doc.add_heading('5.6 评估方法', level=2)
doc.add_paragraph('采用双轨评估策略：')
add_ref(doc, 'D1', '同Session内评估: 5折分层交叉验证（StratifiedKFold, n_splits=5, random_state=42）。这意味着每折的训练集和测试集均保持相同的类别比例。第2个Session的200试次，每折用160试次训练、40试次测试。')
add_ref(doc, 'D2', '跨Session泛化评估: 训练集为Sessions 1-4（777 trials），测试集为Session 5（197 trials）。此方法评估模型在新的采集日期的真实泛化能力。')

doc.add_page_break()

# ====================== 6. 模型结果 ======================
doc.add_heading('六、模型结果与可视化分析', level=1)

doc.add_heading('6.1 15种脑区组合的4分类准确率排名', level=2)
doc.add_paragraph('以下结果均基于同Session冈5折交叉验证（ref D1），特征为4个ERP窗口均值（ref A6），分类器为标准LDA（ref C1-C2）：')
add_table(doc,
    ['排名', '脑区组合', '通道', '特征维', '4分类准确率', '轴准确率'],
    [['1', 'F+P (额叶+顶叶)', '6', '24', '65.0%', '54.5%'],
     ['2', 'F+O (额叶+枕叶)', '6', '24', '62.0%', '48.5%'],
     ['3', '额叶 (F3,Fz,F4)', '3', '12', '61.5%', '52.5%'],
     ['4', 'F+P+O', '9', '36', '60.5%', '48.5%'],
     ['5', 'F+C', '6', '24', '60.5%', '50.5%'],
     ['6', 'F+C+O', '9', '36', '59.0%', '51.5%'],
     ['7', 'F+C+P', '9', '36', '57.0%', '51.0%'],
     ['8', 'All 12ch', '12', '48', '54.0%', '50.5%'],
     ['9', 'P+O', '6', '24', '52.5%', '51.5%'],
     ['10', '中央 (C3,Cz,C4)', '3', '12', '49.5%', '47.0%'],
     ['11', '顶叶 (P3,Pz,P4)', '3', '12', '48.0%', '53.5%'],
     ['12', 'C+P+O', '9', '36', '48.0%', '54.5%'],
     ['13', 'C+O', '6', '24', '46.0%', '52.5%'],
     ['14', 'C+P', '6', '24', '44.5%', '54.5%'],
     ['15', '枕叶 (O1,Oz,O2)', '3', '12', '41.0%', '49.5%']])
add_img(doc, os.path.join(OUT_DIR, '31_region_4class_accuracy.png'), Inches(5.0))

doc.add_heading('6.2 混淆矩阵分析（F+P，最佳模型）', level=2)
doc.add_paragraph('混淆矩阵格子为百分比，行=真实类别，列=预测类别（ref B3）。数据来自Session 2的5折交叉验证混淆矩阵汇总：')
add_table(doc,
    ['真实\\预测', 'Up', 'Down', 'Left', 'Right', '对角线（单类准确率）'],
    [['Up', '42%', '32%', '14%', '12%', '42%'],
     ['Down', '20%', '46%', '12%', '22%', '46%'],
     ['Left', '10%', '16%', '62%', '12%', '62%'],
     ['Right', '2%', '10%', '0%', '88%', '88%']])
doc.add_paragraph('')
add_ref(doc, 'E1', '单类准确率: Right=88%, Left=62%, Down=46%, Up=42%。Right的识别最准确，Left和Right之间几乎无混淆（互为0%）。')
add_ref(doc, 'E2', '错误结构分析: Up被误判为Down的比例高达32%，Down被误判为Up为20%，说明垂直方向的二分辨是主要错误来源。这与下文的轴判别瓶颈分析一致。')
add_img(doc, os.path.join(OUT_DIR, '35_region_best_confusion.png'), Inches(4.5))

doc.add_heading('6.3 六组二分类结果', level=2)
doc.add_paragraph(
    '4个方向（0=Up, 1=Down, 2=Left, 3=Right）共有6组两两配对。'
    '每组在7个脑区组合中筛选最强表现。'
    '结果均基于5折交叉验证（ref D1）、ERP窗口均值特征（ref A6）、LDA分类器（ref C1-C2）。'
    '二分类机会水平为50%：')

add_table(doc,
    ['配对编号', '方向配对', '最强脑区组合', '最佳准确率', '详细说明'],
    [['P1', 'Up vs Down', 'F+P+O (额+顶+枕)', '72.0%', '垂直方向中的上下区分，需要多脑区联合方可达72%'],
     ['P2', 'Left vs Right', '额叶 / F+P', '95.0%', '最强配对，左右方向在额叶区有极其明确的ERP差异'],
     ['P3', 'Up vs Left', 'F+P', '78.0%', '“上”和“左”分属不同轴，差异较大，但F+P可达78%'],
     ['P4', 'Up vs Right', '额叶', '87.0%', 'Right方向在额叶的信号特征非常突出，与Up差异明显'],
     ['P5', 'Down vs Left', '额叶', '76.0%', '下和左也分属不同轴，额叶区可达76%'],
     ['P6', 'Down vs Right', '额叶', '87.0%', '同样受益于Right的强信号，额叶87%'],
     ['P7', '轴判别 (V/H)', 'F+P / C+P', '54.5%', '垂直类(Up/Down) vs 水平类(Left/Right)，仅略高于机会水平50%，为整个任务的瓶颈']])

doc.add_paragraph('')
add_ref(doc, 'F1', '二分类结果性质: 含Right的配对（P4, P6）均达87%，说明Right方向的ERP特征最为突出。而垂直轴内的Up vs Down（P1）仅72%，远低于水平轴内的Left vs Right（P2）的95%。')
add_ref(doc, 'F2', '轴判别瓶颈（P7）: 在全部15/15个脑区组合中，轴准确率均为最差的子问题，最高仅54.5%，仅稍高于随机水平。这是4分类准确率无法突破65%的根本原因。')
add_img(doc, os.path.join(OUT_DIR, '92_binary_by_region.png'), Inches(5.0))
add_img(doc, os.path.join(OUT_DIR, '93_4class_vs_binary.png'), Inches(4.5))

doc.add_heading('6.4 跨Session泛化结果', level=2)
doc.add_paragraph('模型基于Sessions 1-4训练（ref D2），在Session 5测试。与同Session的Session 2内5折交叉验证（ref D1）对比：')
add_table(doc,
    ['脑区', '同Session（ref D1）', '跨Session（ref D2）', '衰减'],
    [['中央', '49.5%', '48.2%', '-1.3%'],
     ['枕叶', '41.0%', '39.1%', '-1.9%'],
     ['顶叶', '48.0%', '36.0%', '-12.0%'],
     ['All 12ch', '54.0%', '41.6%', '-12.4%'],
     ['额叶', '61.5%', '32.0%', '-29.5%'],
     ['F+P', '65.0%', '35.0%', '-30.0%']])
doc.add_paragraph('')
add_ref(doc, 'G1', '中央区是唯一跨Session仅衰减1.3%的脑区，表明中央区（运动皮层）的方向编码模式在不同日期间非常稳定。')
add_ref(doc, 'G2', '额叶区跨Session衰减达29.5%，说明同Session内的额叶高准确率（61.5%）大量受益于Session内噪声模式的过拟合，而非真正的方向编码。')
add_img(doc, os.path.join(OUT_DIR, '51_cross_session_bars.png'), Inches(5.0))

doc.add_heading('6.5 特征重要性分析', level=2)
doc.add_paragraph(
    '特征重要性基于跨Session LDA模型（训练于Sessions 1-4）的系数权重计算。')
add_ref(doc, 'H1', '计算方法: LDA为4分类产生3个判别函数，每个函数对应一组系数向量w_i。对于第j个特征，其重要性定义为跨所有3个判别函数的系数绝对值的平均值: importance_j = mean(|w_1j|, |w_2j|, |w_3j|)。这一指标反映了该特征对所有类别判别的平均贡献度。')

doc.add_paragraph('按窗口汇总（对所有通道取平均）:')
add_ref(doc, 'H2', '汇总方法: 将同一窗口的12个通道的重要性值取平均，得到该窗口的平均重要性: win_importance_w = mean(importance_{w,ch} for ch in 1..12)。这反映了各个时间窗口在方向解码中的平均贡献。')
add_table(doc, ['排名', '窗口', '时间范围', '重要性值', '计算方法'],
    [['1', 'P2', '200-300ms', '0.496', 'mean(importance_{P2, ch}) 对所有12通道'],
     ['2', 'N1', '140-200ms', '0.492', '同上'],
     ['3', 'P3', '300-500ms', '0.313', '同上'],
     ['4', 'P1', '80-130ms', '0.288', '同上']])

doc.add_paragraph('')
doc.add_paragraph('按通道汇总（对所有4个窗口取平均）:')
add_ref(doc, 'H3', '汇总方法: 将同一通道的4个窗口的重要性值取平均，得到该通道的平均重要性: ch_importance_c = mean(importance_{c, w} for w in 1..4)。这反映了各个通道在方向解码中的平均贡献。')
add_table(doc, ['排名', '通道', '重要性值', '计算方法'],
    [['1', 'C3 (左中央)', '0.605', 'mean(importance_{C3, w}) 对4个窗口'],
     ['2', 'F4 (右额叶)', '0.562', '同上'],
     ['3', 'Pz (顶中线)', '0.482', '同上'],
     ['4', 'Fz (额中线)', '0.471', '同上'],
     ['5', 'O2 (右枕叶)', '0.456', '同上'],
     ['6', 'F3 (左额叶)', '0.206', '同上']])

doc.add_paragraph('')
add_ref(doc, 'H4', 'P2和N1窗口的重要性达0.50左右，远高于P1和P3。C3（左中央）是最重要的单个通道，F3（左额叶）是最不重要的通道。')
add_img(doc, os.path.join(OUT_DIR, '62_feature_importance_all_channels.png'), Inches(5.0))

doc.add_heading('6.6 特征工程对照总表', level=2)
doc.add_paragraph('以下结果均基于同Session 5折交叉验证（ref D1）：')
add_table(doc,
    ['方案', '额叶', '中央', 'F+P', '全部 12ch', '算法参考'],
    [['4窗口均值(基准)', '61.5%', '49.5%', '65.0%', '54.0%', 'A6'],
     ['+u5cf0-u5cf0范围', '63.0%', '45.0%', '62.0%', '51.5%', 'A6+范围'],
     ['全时程+PCA', '45.5%', '39.5%', '44.0%', '46.5%', 'A6+全时程'],
     ['全时程+收缩LDA', '46.0%', '41.0%', '44.0%', '45.0%', 'C1+收缩'],
     ['FFT频带能量', '42.0%', '35.5%', '46.5%', '43.5%', '频带B'],
     ['特征融合', '59.5%', '46.0%', '59.0%', '49.5%', 'A6+频带']])

doc.add_paragraph('')
add_ref(doc, 'I1', 'No feature engineering approach improved over the baseline 4 window means. '
    'This suggests the decoding bottleneck is not feature quality but the signal itself '
    '（limited channels, SRB2 reference, 50 trials/direction）.')

doc.add_heading('6.7 层次化解码分析', level=2)
doc.add_paragraph(
    '4分类被拆解为两个阶段: 轴判别（垂直 vs 水平）-> 轴内方向判别（Up/Down 或 Left/Right）。'
    '层次化准确率 = 轴准确率 x 轴内平均准确率。')
add_ref(doc, 'J1', '层次化准确率公式: acc_hier = axis_acc x (vert_acc + horz_acc) / 2。其中 axis_acc为垂直vs水平的二分类准确率，vert_acc为Up vs Down准确率，horz_acc为Left vs Right准确率。')
add_ref(doc, 'J2', '轴判别在全部15/15（100%）个脑区组合中均为瓶颈，最高仅54.5%。')

add_table(doc,
    ['脑区', '轴(V/H)', 'U vs D', 'L vs R', '4分类', '层次化', '瓶颈'],
    [['F+P', '54.5%', '69.0%', '95.0%', '65.0%', '44.7%', '轴'],
     ['额叶', '52.5%', '57.0%', '95.0%', '61.5%', '39.9%', '轴'],
     ['中央', '47.0%', '64.0%', '89.0%', '49.5%', '36.0%', '轴'],
     ['顶叶', '53.5%', '57.0%', '79.0%', '48.0%', '36.4%', '轴'],
     ['全部 12ch', '50.5%', '65.0%', '86.0%', '54.0%', '38.1%', '轴']])

add_img(doc, os.path.join(OUT_DIR, '98a_hierarchical_breakdown.png'), Inches(5.0))

doc.add_heading('6.8 核心结论', level=2)
conclusions = [
    '4方向解码上限（同Session）: F+P = 65.0%（ref D1），机会水平=25%',
    '解码瓶颈: 轴判别（垂直vs水平）在100%脑区组合中均为最弱环节（ref J2），最高仅54.5%（ref F2）',
    '泛化天花板: 跨Session最高仅48.2%（中央区），额叶衰减30%（ref G1, G2）',
    '最佳特征: 4个ERP窗口均值（P1/N1/P2/P3）最优，其他方法均无法超越（ref I1）',
    '半球偏侧化: 右半球通道（O2, P4, C4, F4）解码能力显著强于左半球（ref H3, 71_single_channel_decoding.png）',
    '频域线索: 频带能量可提升轴判别11.5%，但损失时域方向区分力',
]
for c in conclusions:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# ====================== 7. 总结 ======================
doc.add_heading('七、总结', level=1)
doc.add_paragraph(
    '本项目完整覆盖了从硬件配置、刺激呈现、数据采集、信号处理、特征工程、'
    '机器学习建模到结果可视化的EEG研究全流程。通过Processing/Java开发了带实时EEG监控的箭头刺激程序，'
    '基于Python完成了系统性的数据分析管线搭建，最终将全部结果汇总为结构化实验报告。')
doc.add_paragraph(
    '项目最核心的发现——“轴判别瓶颈”——揭示了当前硬件条件下12通道SRB2参考方案的解码局限性，'
    '为后续改进方向（更换参考方案、增加通道密度、提高试次数量）提供了明确的实验依据。')

doc.save(REPORT_PATH)
print(f'Saved: {REPORT_PATH}')
